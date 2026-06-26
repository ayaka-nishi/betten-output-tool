#!/usr/bin/env python3
"""
3号対象返礼品一覧表 → 事業者別証明書（yousiki）生成 / 回収・別添1生成スクリプト

━━ 使い方 ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  【初回】3号対象返礼品.xlsx → まとめ + 送付用証明書を生成
    python3 証明書作成.py

  【更新】証明書まとめ.xlsx を編集した内容を個別ファイルに反映
    python3 証明書作成.py
    ※ 証明書まとめ.xlsx があれば自動的にマスタとして使用

  【回収→別添1生成】事業者から返ってきたファイルを集計して別添1を出力
    1. 返送されたファイルを「証明書回収」フォルダに入れる
    2. python3 証明書作成.py --collect

  【リセット】3号対象返礼品.xlsx からまとめを作り直す
    python3 証明書作成.py --reset

━━ フォルダ構成 ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  証明書出力/                   ← 送付用ファイル + マスタ
    証明書まとめ.xlsx            ← マスタ管理ファイル
    証明書_{事業者名}.xlsx       ← 各事業者に送付するファイル

  証明書回収/                   ← 事業者から返送されたファイルをここに入れる
    証明書_{事業者名}.xlsx

  別添1出力/                    ← --collect 後に生成される最終提出用ファイル
    別添1_{事業者名}.docx        ← Word（事業者ごとに1ファイル）
    別添2.xlsx                   ← Excel（全社まとめ・公式様式）
"""

import copy
import datetime
import os
import re
import sys
from collections import defaultdict

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.worksheet.properties import PageSetupProperties
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers as xl_numbers
from openpyxl.utils import get_column_letter

# ── 設定 ─────────────────────────────────────────────────────────────────────
BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
LIST_FILE     = os.path.join(BASE_DIR, "3号対象返礼品.xlsx")
TEMPLATE_FILE = os.path.join(BASE_DIR, "yousiki.xlsx")
OUTPUT_DIR    = os.path.join(BASE_DIR, "証明書出力")
COLLECT_DIR   = os.path.join(BASE_DIR, "証明書回収")
BETTEN1_DIR   = os.path.join(BASE_DIR, "別添1出力")
SUMMARY_FILE  = os.path.join(OUTPUT_DIR, "証明書まとめ.xlsx")
MAIL_DIR      = os.path.join(OUTPUT_DIR, "郵送用")

# 自治体情報（単独実行時のデフォルト。横展開先では下記MUNICIPALITIESから選択する）
LOCAL_GOV      = "石巻市"
LOCAL_GOV_HEAD = "石巻市長　齋藤　正美"   # 市長名が変わった場合はここを更新
PLACEHOLDER    = "【●●】"

# 横展開対応：選択可能な自治体一覧（都道府県, 市区町村）
MUNICIPALITIES = [
    ("宮城県", "石巻市"),
    ("富山県", "氷見市"),
    ("富山県", "滑川市"),
    ("大分県", "玖珠町"),
    ("北海道", "安平町"),
]


def municipality_label(pref: str, city: str) -> str:
    """プルダウン表示用の文字列（都道府県＋市区町村）を返す"""
    return f"{pref}{city}"


def build_gov_names(pref: str, city: str, mayor_name: str):
    """選択した自治体と市長名から (gov_name, gov_head, area_default) を組み立てる
    市区町村名の末尾（市/町/村）から「市長」「町長」「村長」を自動判定する。
    市長名が未入力の場合は「○○市長　殿」のように氏名部分を省略する"""
    gov_name   = city
    mayor_name = (mayor_name or "").strip()
    gov_head   = f"{city}長　{mayor_name}" if mayor_name else f"{city}長"
    area_default = f"{pref}{city}"
    return gov_name, gov_head, area_default

# リスト列番号（0-based）
COL_COMPANY = 8   # I列: 事業者
COL_PRODUCT = 7   # H列: 商品名
COL_PRICE   = 10  # K列: 提供価格（税込）
COL_TYPE    = 12  # M列: 類型（3のみ対象）
# ─────────────────────────────────────────────────────────────────────────────

DATA_START    = 7   # データ行開始行番号
TEMPLATE_ROWS = 16  # テンプレートのデータ行数（行7〜22）
FOOTER_ROW    = 23  # ※1 注釈行の行番号


def safe_filename(name: str) -> str:
    """ファイル名に使えない文字を '_' に置換"""
    return re.sub(r'[\\/:*?"<>|]', '_', name)


def _today_reiwa() -> str:
    """出力日を「令和○年○月○日」形式で返す"""
    today = datetime.date.today()
    reiwa_year = today.year - 2018
    return f"令和{reiwa_year}年{today.month}月{today.day}日"


# ══════════════════════════════════════════════════════════════════════════════
# 読み込み
# ══════════════════════════════════════════════════════════════════════════════

def read_products(filepath: str, area_default: str = "宮城県石巻市") -> dict:
    """3号対象返礼品.xlsx を読み込み、類型3のみ抽出して {事業者名: [商品リスト]} を返す"""
    wb = load_workbook(filepath, data_only=True)
    ws = wb.active

    companies: dict[str, list] = defaultdict(list)
    for row in ws.iter_rows(min_row=2, values_only=True):
        # 類型が 3（数値）または '3'（文字列）のもののみ対象
        type_val = row[COL_TYPE]
        if type_val != 3 and str(type_val).strip() != "3":
            continue
        company = row[COL_COMPANY]
        name    = row[COL_PRODUCT]
        price   = row[COL_PRICE]
        if not company or not name:
            continue
        companies[str(company).strip()].append({
            "name":   str(name).strip(),
            "price":  price,
            "cost_b": None,
            "area":   area_default,
            "retail": None,
        })
    return dict(companies)


def read_from_summary(filepath: str, area_default: str = "宮城県石巻市") -> dict:
    """
    証明書まとめ.xlsx をマスタとして読み込む。
    区域外費用B・製造加工地・一般販売価格も引き継ぐ。
    """
    wb = load_workbook(filepath, data_only=True)
    ws = wb.active

    companies: dict[str, list] = {}
    current_company = None

    for row in ws.iter_rows(min_row=2, values_only=True):
        cells   = list(row) + [None] * 8
        company = cells[0]   # A: 事業者名
        name    = cells[2]   # C: 返礼品等の名称
        price   = cells[4]   # E: 提供価格A
        cost_b  = cells[5]   # F: 区域外費用B
        area    = cells[6]   # G: 製造・加工地
        retail  = cells[7]   # H: 一般販売価格

        if company and str(company).startswith("合計"):
            break

        if company and str(company).strip():
            current_company = str(company).strip()
            if current_company not in companies:
                companies[current_company] = []

        if not name or not current_company:
            continue

        companies[current_company].append({
            "name":   str(name).strip(),
            "price":  price,
            "cost_b": cost_b,
            "area":   str(area).strip() if area else area_default,
            "retail": retail,
        })

    return companies


def read_from_yousiki(filepath: str, area_default: str = "宮城県石巻市") -> tuple[str, list]:
    """
    返送された証明書ファイル（yousiki形式）から
    事業者名と商品リスト（F・H列含む）を読み込む。
    """
    wb = load_workbook(filepath, data_only=True)
    if "リスト" not in wb.sheetnames:
        return None, []

    ws      = wb["リスト"]
    company = ws["G4"].value
    if not company:
        return None, []

    products = []
    row = DATA_START
    while True:
        name = ws.cell(row, 3).value   # C: 商品名
        if not name:
            break
        products.append({
            "name":   str(name).strip(),
            "price":  ws.cell(row, 5).value,
            "cost_b": ws.cell(row, 6).value,
            "area":   str(ws.cell(row, 7).value).strip()
                      if ws.cell(row, 7).value else area_default,
            "retail": ws.cell(row, 8).value,
        })
        row += 1

    return str(company).strip(), products


# ══════════════════════════════════════════════════════════════════════════════
# 出力：個別証明書 / 別添1（共通フォーマット）
# ══════════════════════════════════════════════════════════════════════════════

def _copy_style(src, dst):
    """セルのスタイルをコピー（フォント・枠線・塗り・配置・書式）"""
    if src.has_style:
        dst.font          = copy.copy(src.font)
        dst.fill          = copy.copy(src.fill)
        dst.border        = copy.copy(src.border)
        dst.alignment     = copy.copy(src.alignment)
        dst.number_format = src.number_format


def _adjust_col_widths(ws, products, extra: int = 0):
    """列幅・行高さ・印刷範囲を調整"""
    # ── 列幅 ─────────────────────────────────────────────────────
    ws.column_dimensions["B"].width = 5    # No.
    ws.column_dimensions["C"].width = 46   # 返礼品等の名称（折り返しOK）
    ws.column_dimensions["D"].width = 12   # 価値の割合
    ws.column_dimensions["E"].width = 12   # 提供価格A
    ws.column_dimensions["F"].width = 29   # 区域外費用B
    ws.column_dimensions["G"].width = 16   # 製造・加工地
    ws.column_dimensions["H"].width = 13   # 一般販売価格

    # ── ヘッダー行の高さ ─────────────────────────────────────────
    ws.row_dimensions[5].height = 65   # 「区域内において生じた価値の割合」等
    ws.row_dimensions[6].height = 75   # 「返礼品等の製造・販売等のために…」等

    # ── 印刷範囲（行挿入分を加算） ───────────────────────────────
    last_print_row = 29 + extra
    ws.print_area = f"A1:I{last_print_row}"

    # ── 印刷設定：横は常に1ページ、縦は商品数で切り替え ─────────
    #   extra==0（16件以下）→ 1ページに収める（自動縮小）
    #   extra>0 （17件以上）→ 縦は自動（複数ページ可）
    #   ※ テンプレートの scale=84 を削除し fitToPage を有効にする
    ws.page_setup.scale       = None   # テンプレートの固定倍率を解除
    ws.sheet_properties.pageSetUpPr = PageSetupProperties(fitToPage=True)
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0 if extra > 0 else 1


def create_yousiki(template_path: str, company_name: str,
                   products: list, output_path: str,
                   mail_version: bool = False,
                   gov_name: str = LOCAL_GOV, gov_head: str = LOCAL_GOV_HEAD,
                   area_default: str = "宮城県石巻市"):
    """テンプレートに事業者データを埋め込んで保存

    mail_version=True のとき：
      - D列（価値の割合）を空欄にする
      - 条件付き書式（赤セル）を削除する
    gov_name/gov_head を指定すると、自治体名・市長名を石巻市以外に切り替えられる
    （横展開対応：F6セルの「区域外で生じた費用」もここで動的に上書きする）
    """
    wb = load_workbook(template_path)
    ws = wb["リスト"]

    ws["B1"] = f"{gov_head}　殿"
    ws["B2"] = (f"　下記の返礼品については、{gov_name}の区域内における工程により、"
                "当該返礼品等の価値の50％以上が生じていることを証明します。")
    ws["G4"] = company_name
    ws["F6"] = f"返礼品等の製造・販売等のために\n{gov_name}の区域外で生じた費用\nB"

    n         = len(products)
    extra     = max(0, n - TEMPLATE_ROWS)
    last_data = DATA_START + TEMPLATE_ROWS - 1

    if extra > 0:
        ws.insert_rows(FOOTER_ROW, extra)
        for ri in range(FOOTER_ROW, FOOTER_ROW + extra):
            for ci in range(2, 9):
                _copy_style(ws.cell(last_data, ci), ws.cell(ri, ci))

    for idx, prod in enumerate(products):
        row = DATA_START + idx

        ws.cell(row, 2).value = idx + 1
        c3 = ws.cell(row, 3)
        c3.value     = prod["name"]
        c3.alignment = Alignment(wrap_text=True, vertical="center")

        if mail_version:
            ws.cell(row, 4).value = None   # 郵送用：空欄（テンプレート残留数式をクリア）
        else:
            ws.cell(row, 4).value         = f'=IF(E{row}="","",(E{row}-F{row})/E{row})'
            ws.cell(row, 4).number_format = "0.0%"

        if prod.get("price") is not None:
            ws.cell(row, 5).value         = prod["price"]
            ws.cell(row, 5).number_format = '#,##0"円"'

        if prod.get("cost_b") is not None:
            ws.cell(row, 6).value         = prod["cost_b"]
            ws.cell(row, 6).number_format = '#,##0"円"'

        ws.cell(row, 7).value = prod.get("area", area_default)

        if prod.get("retail") is not None:
            ws.cell(row, 8).value         = prod["retail"]
            ws.cell(row, 8).number_format = '#,##0"円"'

    for idx in range(n, TEMPLATE_ROWS):
        row = DATA_START + idx
        ws.cell(row, 2).value = None
        if mail_version:
            ws.cell(row, 4).value = None   # 郵送用：空欄
        else:
            ws.cell(row, 4).value         = f'=IF(E{row}="","",(E{row}-F{row})/E{row})'
            ws.cell(row, 4).number_format = "0.0%"

    # 条件付き書式の処理
    last_data_row = DATA_START + TEMPLATE_ROWS - 1 + extra
    new_cf_range  = f"D{DATA_START}:D{last_data_row}"
    for key in list(ws.conditional_formatting._cf_rules.keys()):
        if f"D{DATA_START}" in str(key):
            if mail_version:
                # 郵送用：CF を完全削除（赤セルなし）
                del ws.conditional_formatting._cf_rules[key]
            else:
                # 通常版：データ行末尾まで範囲を更新
                rules = ws.conditional_formatting._cf_rules.pop(key)
                for rule in rules:
                    ws.conditional_formatting.add(new_cf_range, rule)
            break

    # 列幅調整・印刷範囲設定（①）
    _adjust_col_widths(ws, products, extra=extra)

    wb.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
# 出力：まとめファイル
# ══════════════════════════════════════════════════════════════════════════════

def create_summary(companies: dict, output_path: str, area_default: str = "宮城県石巻市"):
    """全事業者・全商品をまとめた管理用 Excel を生成"""
    wb = Workbook()
    ws = wb.active
    ws.title = "証明書まとめ"

    header_font  = Font(name="メイリオ", bold=True, size=10, color="FFFFFF")
    header_fill  = PatternFill("solid", fgColor="2F5496")
    company_fill = PatternFill("solid", fgColor="D9E1F2")
    alt_fill     = PatternFill("solid", fgColor="F2F2F2")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left   = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    right  = Alignment(horizontal="right",  vertical="center")
    thin   = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    COLS = [
        ("事業者名",       22),
        ("No.",             5),
        ("返礼品等の名称", 35),
        ("価値割合",       10),
        ("提供価格A",      14),
        ("区域外費用B",    14),
        ("製造・加工地",   18),
        ("一般販売価格",   14),
    ]

    ws.row_dimensions[1].height = 22
    for ci, (header, width) in enumerate(COLS, start=1):
        c = ws.cell(row=1, column=ci, value=header)
        c.font      = header_font
        c.fill      = header_fill
        c.alignment = center
        c.border    = border
        ws.column_dimensions[get_column_letter(ci)].width = width

    current_row = 2
    for company_idx, (company, products) in enumerate(companies.items()):
        alt = (company_idx % 2 == 1)

        for prod_idx, prod in enumerate(products):
            r = current_row

            c_co = ws.cell(r, 1, value=company if prod_idx == 0 else "")
            c_co.fill      = company_fill
            c_co.font      = Font(name="メイリオ", bold=(prod_idx == 0), size=10)
            c_co.alignment = left
            c_co.border    = border

            c_no = ws.cell(r, 2, value=prod_idx + 1)
            c_no.alignment = center
            c_no.border    = border

            c_name = ws.cell(r, 3, value=prod["name"])
            c_name.alignment = left
            c_name.border    = border

            c_ratio = ws.cell(r, 4)
            c_ratio.value         = f'=IF(E{r}="","",(E{r}-F{r})/E{r})'
            c_ratio.number_format = "0.0%"
            c_ratio.alignment     = center
            c_ratio.border        = border

            c_price = ws.cell(r, 5)
            if prod.get("price") is not None:
                c_price.value = prod["price"]
            c_price.number_format = '#,##0"円"'
            c_price.alignment     = right
            c_price.border        = border

            c_cost = ws.cell(r, 6)
            if prod.get("cost_b") is not None:
                c_cost.value = prod["cost_b"]
            c_cost.number_format = '#,##0"円"'
            c_cost.alignment     = right
            c_cost.border        = border

            c_area = ws.cell(r, 7, value=prod.get("area", area_default))
            c_area.alignment = center
            c_area.border    = border

            c_retail = ws.cell(r, 8)
            if prod.get("retail") is not None:
                c_retail.value = prod["retail"]
            c_retail.number_format = '#,##0"円"'
            c_retail.alignment     = right
            c_retail.border        = border

            if alt:
                for ci in range(2, 9):
                    ws.cell(r, ci).fill = alt_fill

            ws.row_dimensions[r].height = 18
            current_row += 1

    ws.freeze_panes = "A2"

    total_products = sum(len(p) for p in companies.values())
    r = current_row
    for ci in range(1, 9):
        c = ws.cell(r, ci)
        c.fill   = PatternFill("solid", fgColor="2F5496")
        c.font   = Font(name="メイリオ", bold=True, size=10, color="FFFFFF")
        c.border = border
    ws.cell(r, 1).value     = f"合計  {len(companies)}社 / {total_products}件"
    ws.cell(r, 1).alignment = left

    wb.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
# 出力：別添1（Word / 事業者別証明書）
# ══════════════════════════════════════════════════════════════════════════════

def _calc_ratio(price, cost_b):
    """(A-B)/A を % で返す。計算不能な場合は None"""
    try:
        a, b = float(price), float(cost_b)
        return (a - b) / a * 100 if a != 0 else None
    except (TypeError, ValueError):
        return None


def _fmt_yen_w(v):
    if v is None:
        return PLACEHOLDER
    try:
        return f"{int(v):,}"
    except (ValueError, TypeError):
        return PLACEHOLDER


def _fmt_pct_w(v, ndigits=1):
    if v is None:
        return PLACEHOLDER
    try:
        return f"{v:.{ndigits}f}%"
    except (ValueError, TypeError):
        return PLACEHOLDER


def _run_w(para, text, bold=False, size=10.5, color=None):
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "MS Mincho"
    if color:
        r.font.color.rgb = color
    return r


def _para_w(doc, text="", bold=False, size=10.5,
            space_before=0, space_after=0,
            align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    if text:
        _run_w(p, text, bold=bold, size=size)
    return p


def _write_certificate_word(doc, company_name, product, first_page=True,
                             gov_name: str = LOCAL_GOV, gov_head: str = LOCAL_GOV_HEAD):
    """1商品の証明書ページをドキュメントに追記"""
    if not first_page:
        doc.add_page_break()

    A      = product.get("price")
    B      = product.get("cost_b")
    area   = product.get("area") or PLACEHOLDER
    retail = product.get("retail")
    ratio  = _calc_ratio(A, B)

    A_str      = _fmt_yen_w(A)
    B_str      = _fmt_yen_w(B)
    ratio_str  = _fmt_pct_w(ratio)
    retail_str = _fmt_yen_w(retail)

    # タイトル
    _para_w(doc, "別添１", bold=True, size=12, space_after=4,
            align=WD_ALIGN_PARAGRAPH.CENTER)

    # 宛先
    _para_w(doc, f"{gov_head}　殿", size=11, space_after=6)

    # 日付・差出人（右寄せ）
    _para_w(doc, _today_reiwa(), size=10.5,
            align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=1)
    _para_w(doc, company_name, size=11,
            align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

    # 証明文
    cert = (
        f"　「{product['name']}」については、"
        f"{gov_name}の区域内における工程により、"
        f"当該返礼品等の価値の{ratio_str}が生じていることを証明します。"
    )
    _para_w(doc, cert, size=10.5, space_after=6)

    # 算出方法
    _para_w(doc, "上記については、以下の算出方法（該当する算出方法に☑）により算出しています。",
            size=10, space_after=2)
    _para_w(doc, "☑　総務大臣が定める標準的な算出方法", size=10.5, space_after=1)
    _para_w(doc, "　　※標準的な算出方法における算出基礎は以下のとおり。",
            size=9, space_after=1, left_indent=0.5)

    p_A = _para_w(doc, "", size=10.5, space_after=1, left_indent=0.8)
    _run_w(p_A, f"Ａ：当該地方団体による返礼品等の調達費用　　{A_str}　円", size=10.5)

    p_B = _para_w(doc, "", size=10.5, space_after=3, left_indent=0.8)
    _run_w(p_B,
           f"Ｂ：当該返礼品等の製造・販売等のために当該地方団体の区域外で生じた費用"
           f"　　{B_str}　円", size=10.5)

    _para_w(doc, "□　その他の算出方法", size=10.5, space_after=1)
    _para_w(doc, "　　※その他の算出方法とする理由及びその算出方法の詳細は以下のとおり。",
            size=9, space_after=6, left_indent=0.5)

    # 加工地・一般販売価格
    _para_w(doc,
            f"　また、当該返礼品等の製造・加工地（※１）は「{area}」であり、"
            f"一般販売価格は{retail_str}円です（※２）。",
            size=10.5, space_after=6)

    # 同意事項
    _para_w(doc, "なお、当該返礼品等を取り扱うに当たって、下記の事項に同意します。",
            size=10, space_after=2)
    for consent in [
        ("・当該返礼品等については、地場産品基準（平成31年総務省告示第179号第５条）"
         "第８号イ～ハの返礼品等として提出先以外の都道府県又は市区町村が取り扱う場合を除き、"
         "本証明書の提出先以外の都道府県又は市区町村の第３号の返礼品等として取り扱わないこと。"),
        ("・当該返礼品等の付加価値の算出方法等について、地方団体の求めに応じ、"
         "必要な説明や資料提供等を行うこと。"),
    ]:
        _para_w(doc, consent, size=9.5, space_after=2)

    # 記載要領
    _para_w(doc, "記載要領", bold=True, size=9, space_before=6, space_after=2)
    for note in [
        ("※１　返礼品等の製造・加工が行われた場所について、国内の場合は都道府県名及び市区町村名"
         "（例：○○県○○市）、国外の場合は国名を記載すること。"),
        ("※２　当該返礼品等を一般消費者に対して販売する際の通常の価格を記載すること。"
         "なお、当該返礼品等が非売品である場合には、当該返礼品等の類似製品に係る通常の価格を記載すること。"),
    ]:
        _para_w(doc, note, size=8.5, space_after=1)

    # 未入力項目の案内（赤字）
    unfilled = [s for s in (ratio_str, A_str, B_str, retail_str) if PLACEHOLDER in s]
    if unfilled:
        p_g = _para_w(doc, "", size=9, space_before=4)
        _run_w(p_g,
               f"【事業者様へ】「{PLACEHOLDER}」の箇所は確認・ご記入の上ご返送ください。",
               size=9, color=RGBColor(0xC0, 0x00, 0x00))


def create_betten1_word(company_name: str, products: list, output_path: str,
                         gov_name: str = LOCAL_GOV, gov_head: str = LOCAL_GOV_HEAD):
    """別添1 Word ファイル（1社分・商品ごとに1ページ）を生成"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)
    doc.styles["Normal"].font.name = "MS Mincho"
    doc.styles["Normal"].font.size = Pt(10.5)

    for i, prod in enumerate(products):
        _write_certificate_word(doc, company_name, prod, first_page=(i == 0),
                                 gov_name=gov_name, gov_head=gov_head)

    doc.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
# 出力：別添2（Excel / 全社一覧・公式様式）
# ══════════════════════════════════════════════════════════════════════════════

def create_betten2_excel(companies: dict, output_path: str, gov_name: str = LOCAL_GOV):
    """別添2 Excel（全社まとめ・総務省公式様式）を生成"""
    wb = Workbook()
    ws = wb.active
    ws.title = "別添2"

    N_COLS = 8   # A〜H

    # ── スタイル定義 ─────────────────────────────────────────────
    box_font   = Font(name="MS Mincho", bold=True, size=11)
    title_font = Font(name="MS Mincho", bold=True, size=14)
    label_font = Font(name="MS Mincho",             size=10.5)
    h_font     = Font(name="MS Mincho", bold=True, size=10)
    data_font  = Font(name="MS Mincho",             size=10)
    yel_fill   = PatternFill("solid", fgColor="FFFF99")
    center  = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left    = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    right   = Alignment(horizontal="right",  vertical="center")
    rt_top  = Alignment(horizontal="right",  vertical="center")
    thin    = Side(style="thin", color="000000")
    border  = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── 列幅 ─────────────────────────────────────────────────────
    widths = [28, 12, 12, 18, 18, 16, 16, 16]
    for ci, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    # ── タイトル部（row1〜3） ─────────────────────────────────────
    ws.row_dimensions[1].height = 26
    c_box = ws.cell(1, 1, value="別添２")
    c_box.font      = box_font
    c_box.alignment = center
    c_box.border    = border

    ws.merge_cells(start_row=1, start_column=2, end_row=1, end_column=N_COLS)
    c_title = ws.cell(1, 2, value="ふるさと納税の返礼品等の区域内において生じた価値の割合に係る一覧表")
    c_title.font      = title_font
    c_title.alignment = Alignment(horizontal="center", vertical="center")

    ws.row_dimensions[2].height = 16
    ws.merge_cells(start_row=2, start_column=5, end_row=2, end_column=N_COLS)
    c_date = ws.cell(2, 5, value=f"公表年月日　{_today_reiwa()}")
    c_date.font      = label_font
    c_date.alignment = rt_top

    ws.row_dimensions[3].height = 16
    ws.merge_cells(start_row=3, start_column=5, end_row=3, end_column=N_COLS)
    c_gov = ws.cell(3, 5, value=f"（地方団体名）　{gov_name}")
    c_gov.font      = label_font
    c_gov.alignment = rt_top

    # ── ヘッダー（3行：5,6,7行目） ───────────────────────────────
    HEADER_TOP, HEADER_MID, HEADER_BOT = 5, 6, 7
    ws.row_dimensions[HEADER_TOP].height = 20
    ws.row_dimensions[HEADER_MID].height = 20
    ws.row_dimensions[HEADER_BOT].height = 32

    def _h(row_s, col_s, row_e, col_e, text):
        ws.merge_cells(start_row=row_s, start_column=col_s, end_row=row_e, end_column=col_e)
        c = ws.cell(row_s, col_s, value=text)
        c.font      = h_font
        c.alignment = center
        for r in range(row_s, row_e + 1):
            for cc in range(col_s, col_e + 1):
                ws.cell(r, cc).border = border
        return c

    _h(HEADER_TOP, 1, HEADER_BOT, 1, "返礼品等の名称")
    _h(HEADER_TOP, 2, HEADER_BOT, 2, "区域内において\n生じた価値の割合\n（％）")
    _h(HEADER_TOP, 3, HEADER_TOP, 5, "区域内において生じた価値の割合の算出方法\n※１")
    _h(HEADER_MID, 3, HEADER_BOT, 3, "標準的な\n算出方法")
    _h(HEADER_MID, 4, HEADER_MID, 5, "その他の\n算出方法")
    _h(HEADER_BOT, 4, HEADER_BOT, 4, "その他の\n算出方法の詳細")
    _h(HEADER_BOT, 5, HEADER_BOT, 5, "その他の\n算出方法とする理由")
    _h(HEADER_TOP, 6, HEADER_BOT, 6, "返礼品等の\n製造・加工地\n※２")
    _h(HEADER_TOP, 7, HEADER_BOT, 7, "地方団体における\n調達費用\n（円）")
    _h(HEADER_TOP, 8, HEADER_BOT, 8, "一般販売価格\n（円）\n※３")

    # ── データ行 ──────────────────────────────────────────────────
    current_row = HEADER_BOT + 1
    for company, products in companies.items():
        for prod in products:
            r = current_row
            A      = prod.get("price")
            B      = prod.get("cost_b")
            ratio  = _calc_ratio(A, B)
            area   = prod.get("area", "") or ""
            retail = prod.get("retail")

            row_vals = [
                (prod["name"], left,   None),
                (ratio / 100 if ratio is not None else None,
                               center, "0.0%"),
                ("○",          center, None),   # 標準的な算出方法
                ("",           left,   None),   # その他の算出方法の詳細
                ("",           left,   None),   # その他の算出方法とする理由
                (area,         center, None),
                (A,            right,  '#,##0"円"'),
                (retail,       right,  '#,##0"円"'),
            ]

            for ci, (val, align, fmt) in enumerate(row_vals, start=1):
                c = ws.cell(r, ci, value=val)
                c.font      = data_font
                c.alignment = align
                c.border    = border
                if fmt:
                    c.number_format = fmt
                if ci == 8 and val is None:   # 一般販売価格：未入力は事業者記入欄として黄色
                    c.fill = yel_fill

            ws.row_dimensions[r].height = 18
            current_row += 1

    # ── 脚注 ─────────────────────────────────────────────────────
    nr = current_row + 1
    notes = [
        ("※１　区域内において生じた価値の割合の算出にあたって、総務大臣が定める標準的な算出方法を用いた場合は"
         "「標準的な算出方法」欄に「○」を記載し、その他の算出方法を用いた場合は「その他の算出方法」欄に「○」を記載した上で、"
         "その算出方法の詳細及びその算出方法とする理由を記載すること。"),
        "なお、区域内において生じた価値の割合の標準的な算出方法は、下記のとおりであること。",
        "　算式",
        "　　（Ａ－Ｂ）／Ａ",
        "　算式の符号",
        "　　Ａ：当該地方団体による返礼品等の調達費用",
        "　　Ｂ：当該返礼品等の製造・販売等のために当該地方団体の区域外で生じた費用",
        ("※２　返礼品等の製造・加工が行われた場所について、国内の場合は都道府県名及び市区町村名"
         "（例：○○県○○市）、国外の場合は国名を記載すること。"),
        ("※３　当該返礼品等を一般消費者に対して販売する際の通常の価格を記載すること。"
         "なお、当該返礼品等が非売品である場合には、当該返礼品等の類似製品に係る通常の価格を記載すること。"),
    ]
    for i, note in enumerate(notes):
        ws.merge_cells(start_row=nr + i, start_column=1, end_row=nr + i, end_column=N_COLS)
        c = ws.cell(nr + i, 1, value=note)
        c.font      = Font(name="MS Mincho", size=8)
        c.alignment = left
        ws.row_dimensions[nr + i].height = 13

    ws.freeze_panes = f"A{HEADER_BOT + 1}"
    wb.save(output_path)


# ══════════════════════════════════════════════════════════════════════════════
# メイン
# ══════════════════════════════════════════════════════════════════════════════

def _print_sep():
    print("─" * 55)


def main():
    reset_mode   = "--reset"   in sys.argv or "-r" in sys.argv
    collect_mode = "--collect" in sys.argv or "-c" in sys.argv

    # ────────────────────────────────────────────────────────────
    # モード①：回収 → まとめ更新 → 別添1生成
    # ────────────────────────────────────────────────────────────
    if collect_mode:
        print(f"\n=== 回収データ集計 → 別添1生成 ===\n")

        if not os.path.exists(SUMMARY_FILE):
            print(f"  ⚠ 証明書まとめ.xlsx が見つかりません")
            print(f"    先に通常実行（引数なし）で初期ファイルを作成してください\n")
            return

        if not os.path.exists(COLLECT_DIR):
            os.makedirs(COLLECT_DIR, exist_ok=True)
            print(f"  ⚠ 証明書回収フォルダを作成しました")
            print(f"    返送されたファイルをここに入れて再実行してください:")
            print(f"    {COLLECT_DIR}\n")
            return

        # 既存まとめからベースデータを読み込む
        base_companies = read_from_summary(SUMMARY_FILE)

        # 回収フォルダのファイルを読み込む
        print("【回収ファイル読み込み】")
        files = sorted([
            f for f in os.listdir(COLLECT_DIR)
            if f.startswith("証明書_") and f.endswith(".xlsx")
        ])

        if not files:
            print(f"  ⚠ 証明書回収フォルダにファイルがありません")
            print(f"    返送されたファイルをこのフォルダに入れてください:\n"
                  f"    {COLLECT_DIR}\n")
            return

        collected = {}
        for fname in files:
            fpath = os.path.join(COLLECT_DIR, fname)
            company, products = read_from_yousiki(fpath)
            if not company:
                print(f"  ⚠ {fname}: 事業者名が読み取れませんでした")
                continue
            n_f = sum(1 for p in products if p["cost_b"] is not None)
            n_h = sum(1 for p in products if p["retail"]  is not None)
            print(f"  ✓ {fname}")
            print(f"      {len(products)}件  |  区域外費用B: {n_f}件入力  |  "
                  f"一般販売価格: {n_h}件入力")
            collected[company] = products

        # まとめ.xlsx に回収データをマージ（回収済みは上書き）
        merged = dict(base_companies)
        for company, products in collected.items():
            merged[company] = products

        # まとめ.xlsx を更新
        create_summary(merged, SUMMARY_FILE)

        # 回収状況サマリー
        _print_sep()
        print(f"\n【回収状況】")
        all_companies = list(merged.keys())
        for company in all_companies:
            if company in collected:
                products = collected[company]
                n_f = sum(1 for p in products if p["cost_b"] is not None)
                n_h = sum(1 for p in products if p["retail"]  is not None)
                total = len(products)
                ok   = (n_f == total and n_h == total)
                mark = "✓ 回収済" if ok else "△ 一部未記入"
                print(f"  {mark}  {company}  ({n_f}/{total}件 B入力, "
                      f"{n_h}/{total}件 H入力)")
            else:
                n = len(merged[company])
                print(f"  ✗ 未回収   {company}  ({n}件)")

        returned_count = len(collected)
        print(f"\n  回収: {returned_count}/{len(all_companies)}社")

        # 別添1（Word）を生成（全社分。未回収社は●●プレースホルダー）
        os.makedirs(BETTEN1_DIR, exist_ok=True)
        _print_sep()
        print(f"\n【別添1 生成（Word）】")
        for company, products in merged.items():
            fname       = f"別添1_{safe_filename(company)}.docx"
            output_path = os.path.join(BETTEN1_DIR, fname)
            create_betten1_word(company, products, output_path)
            status = "完成" if company in collected else "未回収（●●あり）"
            print(f"  ✓ {fname}  [{status}]")

        # 別添2（Excel）を生成（全社まとめ）
        _print_sep()
        print(f"\n【別添2 生成（Excel）】")
        betten2_path = os.path.join(BETTEN1_DIR, "別添2.xlsx")
        create_betten2_excel(merged, betten2_path)
        total_prods = sum(len(p) for p in merged.values())
        print(f"  ✓ 別添2.xlsx  ({len(merged)}社 / {total_prods}件)")

        _print_sep()
        print(f"\n  更新: 証明書まとめ.xlsx")
        print(f"  出力: {BETTEN1_DIR}")
        print("完了\n")
        return

    # ────────────────────────────────────────────────────────────
    # モード②：初回作成 または まとめからの更新
    # ────────────────────────────────────────────────────────────
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    summary_exists = os.path.exists(SUMMARY_FILE)

    if not reset_mode and summary_exists:
        print(f"\n=== 証明書更新（マスタ読み込み）===")
        print(f"  📋 証明書まとめ.xlsx をマスタとして使用します")
        print(f"     ※ テスト3号対象から作り直す場合は --reset を付けて実行\n")
        companies      = read_from_summary(SUMMARY_FILE)
        update_summary = False
    else:
        label = "（--reset）" if reset_mode else ""
        print(f"\n=== 証明書新規作成{label} ===")
        print(f"  📄 3号対象返礼品.xlsx から作成します\n")
        companies      = read_products(LIST_FILE)
        update_summary = True

    print(f"  対象事業者: {len(companies)}社\n")

    # --reset 時は古い証明書_*.xlsx を事前に削除（孤立ファイル防止）
    if update_summary:
        old_files = [
            f for f in os.listdir(OUTPUT_DIR)
            if f.startswith("証明書_") and f.endswith(".xlsx")
        ]
        if old_files:
            for f in old_files:
                os.remove(os.path.join(OUTPUT_DIR, f))
            print(f"  🗑 古い証明書ファイル {len(old_files)}件を削除\n")

    os.makedirs(MAIL_DIR, exist_ok=True)

    for company, products in companies.items():
        safe = safe_filename(company)

        # 通常版（価値の割合あり）
        fname       = f"証明書_{safe}.xlsx"
        output_path = os.path.join(OUTPUT_DIR, fname)
        create_yousiki(TEMPLATE_FILE, company, products, output_path)

        # 郵送用（価値の割合・CF 空欄）
        mail_path = os.path.join(MAIL_DIR, f"証明書_{safe}_郵送用.xlsx")
        create_yousiki(TEMPLATE_FILE, company, products, mail_path, mail_version=True)

        print(f"  ✓ {fname}  ({len(products)}件)")

    if update_summary:
        create_summary(companies, SUMMARY_FILE)
        print(f"\n  ✓ 証明書まとめ.xlsx  （新規作成・マスタ）")
    else:
        print(f"\n  ─ 証明書まとめ.xlsx  （マスタのため変更なし）")

    print(f"  ✓ 郵送用/  （D列空欄・CF解除版 {len(companies)}社）")

    _print_sep()
    print(f"\n  次のステップ:")
    print(f"  1. 証明書出力/ の 証明書_{{事業者名}}.xlsx を各社に送付")
    print(f"  2. 返送されたファイルを「証明書回収」フォルダに入れる")
    print(f"  3. python3 証明書作成.py --collect を実行")
    print(f"\n  出力先: {OUTPUT_DIR}")
    print("完了\n")


if __name__ == "__main__":
    main()
